# write a word document from python test file
import docx

class Document():

    def __init__(self):
        self.doc = docx.Document()
        self.p = ""
        self.default_bold = False
        self.default_italic = False
        self.font = 'Calibri'
        self.font_size = docx.shared.Pt(11)

    def new_paragraph(self,
                      line_spacing=1,
                      space_after=0,
                      ):
        
        # add a new paragraph
        self.p = self.doc.add_paragraph()

        # set paragraph formatting
        self.p.paragraph_format.line_spacing = line_spacing
        self.p.paragraph_format.space_after = space_after

        return self.p

    def write(self,
              text: str,
              paragraph,
              bold: bool,
              italic: bool,
              font: str,
              size):
        
        # add text as run to paragraph
        run = paragraph.add_run(text)

        # format that bitch
        run.bold = bold
        run.italic = italic
        run.font.name = font
        run.font.size = docx.shared.Pt(16)

    def save(self, docname):
        self.doc.save(docname)


def main():
    word_doc = 'C:/Users/kleinj/Downloads/Code/remove_markdown/test_doc.docx'

    doc = Document()
    paragraph1 = doc.new_paragraph()

    text = "Test Document Writing"
    doc.write(text, paragraph1, True, False, 'Calibri', '12')
    doc.save(word_doc)



if __name__ == '__main__':
    main()
        
